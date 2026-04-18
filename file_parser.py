from docx import Document
from zipfile import BadZipFile


def parse_docx(file):
    """
    解析 .docx 简历文件，提取段落和表格文本。
    成功返回纯文本字符串，失败返回「【错误】...」格式字符串。
    """
    try:
        file.seek(0)
        doc = Document(file)

        content = []
        for p in doc.paragraphs:
            if p.text.strip():
                content.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    content.append(" | ".join(cells))

        result = "\n".join(content)
        if not result.strip():
            return "【错误】文件内容为空，请检查简历是否有文字内容"
        return result

    except BadZipFile:
        return "【错误】文件损坏或非标准 .docx 格式"
    except Exception as e:
        return "【解析失败】" + str(e)

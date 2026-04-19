from io import BytesIO


def parse_docx(file):
    try:
        from docx import Document
    except Exception as exc:
        return f"【解析失败】缺少 python-docx：{exc}"

    try:
        if hasattr(file, "getvalue"):
            raw_bytes = file.getvalue()
        elif hasattr(file, "read"):
            raw_bytes = file.read()
        else:
            raw_bytes = file

        stream = BytesIO(raw_bytes)
        doc = Document(stream)

        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))

        text = "\n".join(parts).strip()
        if not text:
            return "【空文档】"

        return text
    except Exception as exc:
        return f"【解析失败】{exc}"

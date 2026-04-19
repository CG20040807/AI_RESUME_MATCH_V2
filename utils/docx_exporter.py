import re
from io import BytesIO
from typing import Dict, Any, List


def _safe_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _join_list(items) -> str:
    if not items:
        return "暂无"
    return "；".join(_safe_text(x) for x in items if _safe_text(x)) or "暂无"


def _append_markdown_like(doc, text: str):
    if not text:
        return

    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", s), style="List Number")
        else:
            doc.add_paragraph(s)


def _setup_doc(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = normal.font.size or None


def _write_job_info(doc, job_title, jd, criteria):
    doc.add_heading("岗位信息", level=1)
    doc.add_paragraph(f"岗位名称：{job_title}")
    doc.add_paragraph("岗位JD：")
    _append_markdown_like(doc, jd)
    doc.add_paragraph("评估标准：")
    _append_markdown_like(doc, criteria)


def _write_ranking_table(doc, ranked):
    doc.add_heading("候选人排名", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = ["排名", "姓名", "总分", "推荐意见", "置信度", "摘要"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header

    for item in ranked:
        row = table.add_row().cells
        row[0].text = str(item.get("rank", ""))
        row[1].text = _safe_text(item.get("name", ""))
        row[2].text = str(item.get("score", 0))
        row[3].text = _safe_text(item.get("recommendation", ""))
        row[4].text = str(item.get("confidence", 0))
        row[5].text = _safe_text(item.get("summary", ""))


def _write_candidate_detail(doc, candidate: Dict[str, Any]):
    doc.add_heading(f"{candidate.get('rank', '')} - {candidate.get('name', '')}", level=2)
    doc.add_paragraph(f"总分：{candidate.get('score', 0)}")
    doc.add_paragraph(f"推荐意见：{candidate.get('recommendation', '')}")
    doc.add_paragraph(f"置信度：{candidate.get('confidence', 0)}")

    if candidate.get("sub_scores"):
        doc.add_paragraph("分项评分：")
        for k, v in candidate["sub_scores"].items():
            doc.add_paragraph(f"{k}：{v if v is not None else '未提取'}", style="List Bullet")

    doc.add_paragraph("摘要：")
    _append_markdown_like(doc, candidate.get("summary", ""))

    doc.add_paragraph("匹配原因：")
    _append_markdown_like(doc, candidate.get("match_reason", ""))

    doc.add_paragraph("优势：")
    for x in candidate.get("strengths", []) or []:
        doc.add_paragraph(_safe_text(x), style="List Bullet")

    doc.add_paragraph("不足：")
    for x in candidate.get("weaknesses", []) or []:
        doc.add_paragraph(_safe_text(x), style="List Bullet")

    doc.add_paragraph("风险：")
    for x in candidate.get("risks", []) or []:
        doc.add_paragraph(_safe_text(x), style="List Bullet")

    doc.add_paragraph("面试问题：")
    for x in candidate.get("questions", []) or []:
        doc.add_paragraph(_safe_text(x), style="List Number")


def export_to_docx(job_title, jd, criteria, ranked, summary_text):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f"缺少 python-docx：{exc}")

    doc = Document()
    _setup_doc(doc)

    doc.add_heading("AI 简历评估报告", level=0)
    doc.add_paragraph(f"生成时间：{__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"岗位名称：{job_title}")

    _write_job_info(doc, job_title, jd, criteria)
    doc.add_heading("综合总结", level=1)
    _append_markdown_like(doc, summary_text)

    _write_ranking_table(doc, ranked)

    doc.add_heading("候选人详情", level=1)
    for item in ranked:
        _write_candidate_detail(doc, item)
        doc.add_paragraph("")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def export_single_to_docx(job_title, jd, criteria, candidate):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f"缺少 python-docx：{exc}")

    doc = Document()
    _setup_doc(doc)

    doc.add_heading("候选人评估报告", level=0)
    doc.add_paragraph(f"生成时间：{__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"岗位名称：{job_title}")
    _write_job_info(doc, job_title, jd, criteria)
    _write_candidate_detail(doc, candidate)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
